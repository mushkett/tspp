from sqlalchemy.exc import DataError
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from model import Flight
import PySimpleGUI as Sg

from docx import Document


def create_session():
    host = "ec2-54-165-178-178.compute-1.amazonaws.com"
    db = "dbvhl0gmdq014r"
    user = "glygllyaluruqa"
    port = "5432"
    password = "dd27e4dc8a6c33f7842bfea2c025ebf2536b93d235c9239e13ffe23cca3e439b"

    engine = create_engine(f"postgresql+psycopg2://{user}:{password}@{host}:{port}/{db}")
    Session = sessionmaker(bind=engine)
    session = Session()
    return session


def get_flights(plane_number, hours):
    session = create_session()
    result = []
    try:
        if hours and plane_number:

            result = session.query(Flight)\
                .filter(Flight.hours < int(hours))\
                .filter(Flight.plain_number == plane_number)\
                .order_by(Flight.id)\
                .all()

        elif hours:
            result = session.query(Flight)\
                .filter(Flight.hours >= int(hours)) \
                .order_by(Flight.id)\
                .all()
        elif plane_number:
            result = session.query(Flight)\
                .filter(Flight.plain_number == plane_number) \
                .order_by(Flight.id) \
                .all()
        else:
            result = session.query(Flight).order_by(Flight.id).all()

    except DataError:
        session.rollback()
        Sg.Popup('Виникла помилка! Перевірте правильність введених даних!', title="Error")
    else:
        if result:
            list_of_flights = []

            for i in result:
                temp_list = []
                i_dict = i.__dict__
                temp_list.append(i_dict['id'])
                temp_list.append(i_dict['plain_number'])
                temp_list.append(i_dict['hours'])
                temp_list.append(i_dict['route_number'])
                list_of_flights.append(temp_list)
            session.close()
            return list_of_flights
        else:
            Sg.Popup('Нічого не знайдено!', title='Error')


def get_plane_list():
    session = create_session()
    result = sorted(set([i[0] for i in session.query(Flight.plain_number).all()]))
    return result


def save_to_docx(list_of_flights):
    document = Document()
    document.add_heading("Flights", 1)
    if list_of_flights:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = "Номер літака"
        hdr_cells[2].text = "Кількість льотних годин"
        hdr_cells[3].text = "Номер рейсу"

        for item in list_of_flights:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item[0])
            row_cells[1].text = item[1]
            row_cells[2].text = str(item[2])
            row_cells[3].text = item[3]
        document.save("Flights.docx")
    else:
        return Sg.Popup("Flights not found")

    return Sg.popup("Document was created successfully!", title="Success!")


def authorization():
    is_authorized = False
    log = 'admin'
    passw = 'admin'

    login = Sg.popup_get_text("Input login", "Authorization")
    password = Sg.popup_get_text("Input password", "Authorization", password_char='*')
    if log == login and passw == password:
        is_authorized = True
        Sg.popup("You are authorized now!", title="Success!")
    else:
        Sg.Popup("Incorrect login or password", title="Error")
    return is_authorized


def add_flight(flight_id, plain_number, hours, route_number):
    session = create_session()
    if flight_id:
        flight = session.query(Flight).filter(Flight.id == int(flight_id))
        record = flight.one()
        record.plain_number = plain_number
        record.hours = int(hours)
        record.route_number = route_number
        Sg.Popup('Рейс успішно оновлено', title='Success')

    else:
        new_flight = Flight(
            plain_number=plain_number,
            hours=hours,
            route_number=route_number
        )
        session.add(new_flight)
        Sg.Popup('Рейс успішно додано', title='Success')

    session.commit()
