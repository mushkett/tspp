import PySimpleGUI as Sg

from docx import Document


def save_to_docx(list_of_flights):
    document = Document()
    document.add_heading("Flights", 1)
    if list_of_flights:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = "Номер літака"
        hdr_cells[2].text = "Кількість льотних годин"
        hdr_cells[3].text = "Номер рейсу"

        for item in list_of_flights:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item[0])
            row_cells[1].text = item[1]
            row_cells[2].text = str(item[2])
            row_cells[3].text = item[3]
        document.save("Flights.docx")
    else:
        return Sg.Popup("Flights not found")

    return Sg.popup("Document was created successfully!", title="Success!")
